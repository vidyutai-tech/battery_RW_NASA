from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION

OUT = "/Users/himasoni_123/Documents/battery_RW_NASA/deliverables/hybrid_degradation_methodology_results_discussion_draft.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
LIGHT = "F4F6F9"
GRAY = "5B6573"
GOLD = "7A5A00"

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def set_fixed_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(width))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_run(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic

def add_paragraph(doc, text="", style=None, size=11, bold=False, color=INK, italic=False):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color, italic=italic)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(text), size=10.5)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(text), size=10.5)
    return p

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, headers):
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run(p.add_run(text), size=9.5, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(str(text)), size=9.3)
    set_fixed_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table

def add_note(doc, label, text, color=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_fixed_table(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7E6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(label + ": "), size=10.2, bold=True, color=color)
    set_run(p.add_run(text), size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_code_prompt(doc, title, prompt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(title), size=12, bold=True, color=DARK_BLUE)
    table = doc.add_table(rows=1, cols=1)
    set_fixed_table(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(prompt)
    set_run(r, size=9.2, color="263238")

def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.33

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("Hybrid Degradation-Aware Charging | Research Draft"), size=8.5, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("Hima Soni | Working draft for supervisory review"), size=8.5, color=GRAY)

def main():
    doc = Document()
    setup(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run("METHODOLOGY AND RESULTS-DISCUSSION DRAFT"), size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run("Hybrid Degradation-Aware, Energy-Constrained Charging Optimization Using a Transfer-Learned Battery Digital Twin"), size=22, bold=True, color="0B2545")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    set_run(p.add_run("Working research draft | 24 July 2026"), size=11, color=GRAY)
    add_note(doc, "Purpose", "This editable draft consolidates the current physical modelling rationale, completed results, implementation priorities, and two prompts for efficient follow-up work. Replace bracketed placeholders only after the corresponding experiment is run.", color=BLUE)
    add_paragraph(doc, "Scope and evidence status", style="Heading 1")
    add_paragraph(doc, "The completed work consists of a transfer-learned battery digital twin, energy-constrained charging simulation, and profile optimization. Earlier results used a temperature-time reward with random search. The current implementation introduces a hybrid calendar-cyclic capacity-loss objective and GP Bayesian optimization. Reinforcement learning is a planned next stage and must not be described as a completed experimental result until trained and evaluated under the same constraints.")
    add_note(doc, "Scientific caution", "The semi-empirical aging coefficients are drawn from literature sources with chemistry- and test-specific conditions. Until calibrated to the NASA RW cells, the reported Q_loss values should be called a relative capacity-loss index rather than an absolute percentage of capacity fade.")

    add_paragraph(doc, "1. Research objective and hypotheses", style="Heading 1")
    add_paragraph(doc, "The objective is to determine a feasible charging-current trajectory that delivers a specified energy requirement within an available charging interval while minimizing modelled calendar and cycling degradation. The battery state is propagated through a transfer-learned digital twin and a Coulomb-counted SoC state update.")
    add_bullet(doc, "H1: Under a fixed energy requirement and identical safety limits, the hybrid objective selects profiles with a more interpretable time-versus-degradation trade-off than the earlier temperature-time reward.")
    add_bullet(doc, "H2: At limited evaluation budgets, GP Bayesian optimization reaches lower feasible loss than random search for at least some profile families, particularly pulsed charging with higher-dimensional parameters.")
    add_bullet(doc, "H3: The preferred profile changes with initial SoC, required energy, available time, ambient temperature, and cell age; therefore a single lifetime profile is not sufficient for daily-use charging decisions.")

    add_paragraph(doc, "2. Methodology", style="Heading 1")
    add_paragraph(doc, "2.1 Battery digital twin and state propagation", style="Heading 2")
    add_paragraph(doc, "A Transformer-decoder battery digital twin predicts terminal voltage V(t) and cell temperature T(t) from relative age, initial voltage, initial temperature, and the commanded charging-current sequence. The source model is trained on RW9 and fine-tuned on RW10-RW12. State of charge is not inferred from loaded terminal voltage; instead, it is updated through Coulomb counting using the age-dependent rated capacity and an OCV-SoC initialization.")
    add_paragraph(doc, "2.2 Energy-constrained charging problem", style="Heading 2")
    add_paragraph(doc, "For each charging session, the energy constraint is E_delivered >= E_required. E_delivered is obtained by integrating -V(t) I(t) over time. The decision variables are the parameters of CCCV, two-step, three-step, and pulsed charging families. Feasibility requires the energy target (or SoC target, where used) to be met within the maximum duration while respecting voltage and current limits.")
    add_paragraph(doc, "2.3 Hybrid degradation model", style="Heading 2")
    add_paragraph(doc, "The hybrid model separates calendar and cyclic stress rather than treating temperature only as a shaped reward. The total relative capacity-loss index is Q_loss,total = Q_calendar + Q_cyclic. The primary calendar model is the SoC-temperature-time Arrhenius relation (Eq. 2 in Sinha et al.). The primary cyclic model is the throughput-temperature-C-rate power-law relation (Eq. 7 and Table 7).")
    add_table(doc,
        ["Literature relation", "Physical stress represented", "Use in this work"],
        [
            ["Eq. 2", "Calendar loss: mean SoC, T, time; Arrhenius dependence", "Primary Q_calendar model"],
            ["Eq. 3", "Alternative empirical calendar model", "Sensitivity analysis; do not add to Eq. 2"],
            ["Eq. 4", "Equivalent full cycles, mean SoC and SoC swing", "Report EFC, mean SoC and DeltaSoC; future stress multiplier"],
            ["Eq. 5", "Temperature-time calendar loss", "Alternative calendar sensitivity model"],
            ["Eq. 6", "Generic Arrhenius aging rate", "Physical rationale, not a separate additive loss"],
            ["Eq. 7 / Table 7", "C-rate, T and Ah-throughput cyclic loss", "Primary Q_cyclic model"],
            ["Eq. 8", "Continuous current-dependent cyclic model", "Future alternative after chemistry-specific calibration"],
        ], [1800, 3400, 4160])
    add_note(doc, "Model rule", "Equations 2, 3 and 5 are alternatives for calendar aging. Equations 7/Table 7 and 8 are alternatives for cyclic aging. They should not be summed together without an experimentally justified coupling model.")
    add_paragraph(doc, "2.4 Optimization objective", style="Heading 2")
    add_paragraph(doc, "The session reward balances useful charge delivery, modelled degradation, and time: R = w_soc DeltaSoC - w_qloss(Q_calendar + Q_cyclic) - w_time t_h^z. The scalar loss is -R plus soft penalties for energy/SoC shortfall and voltage overshoot. For physical reporting, reward is secondary; energy, C-rate, temperature and each degradation component are reported separately.")
    add_paragraph(doc, "2.5 Optimization and comparison protocol", style="Heading 2")
    add_paragraph(doc, "Each family is optimized independently by Gaussian-process Bayesian optimization (PI acquisition in the completed RW9 experiments) and compared with random search under the same start state, energy/SoC requirement, current bounds, voltage ceiling, maximum duration, objective weights, and evaluation budget. A fair comparison evaluates the best feasible candidate rather than merely the best unconstrained score.")

    add_paragraph(doc, "3. Results and discussion", style="Heading 1")
    add_paragraph(doc, "3.1 Transfer-learning evidence", style="Heading 2")
    add_paragraph(doc, "The two-stage fine-tuning procedure improves target-cell voltage prediction as more target-cell adaptation data are used. At 60% adaptation data, held-out voltage RMSE was 0.0323 V for RW10, 0.0483 V for RW11, and 0.0639 V for RW12. These results support use of target-specific twins for charging optimization, while the larger RW12 error motivates caution in extrapolating profile rankings across cells.")
    for text in [
        "RW10, 60% adaptation: voltage RMSE = 0.0323 V; temperature RMSE = 0.2842 deg C.",
        "RW11, 60% adaptation: voltage RMSE = 0.0483 V; temperature RMSE = 0.1517 deg C.",
        "RW12, 60% adaptation: voltage RMSE = 0.0639 V; temperature RMSE = 0.2636 deg C.",
    ]:
        p = add_paragraph(doc, text, size=10.5)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(3)
    add_paragraph(doc, "3.2 Hybrid-objective GP-BO versus random search on RW9", style="Heading 2")
    add_paragraph(doc, "For the completed RW9 comparison, the start state was SoC 20%, target SoC 80%, T0 = 24 deg C, age = 0, and the maximum duration was 150 min. At 40 evaluations per profile family, GP-BO selected a feasible pulsed profile with loss -0.3516, total reward 0.4121, duration 60.5 min, and relative total loss index 0.0937. The best random-search result was a three-step profile with loss -0.3450. Thus, GP-BO gave the best overall feasible solution at the limited 40-evaluation budget, while the family-level results show that the current three-step GP configuration remains a meaningful optimization limitation.")
    add_table(doc, ["Method / family", "Budget", "Best loss", "Duration (min)", "Q_loss index", "Interpretation"], [
        ["GP-BO / pulsed", "40", "-0.3516", "60.5", "0.0937", "Best overall feasible result"],
        ["Random / 3-step", "40", "-0.3450", "63.0", "0.0929", "Strong random baseline"],
        ["GP-BO / pulsed", "80", "-0.3454", "63.5", "0.0936", "Near-tie at higher budget"],
        ["Random / CCCV", "80", "-0.3452", "63.5", "0.0923", "Near-tie at higher budget"],
    ], [2500, 900, 1150, 1350, 1300, 2160])
    add_paragraph(doc, "Discussion: the 40-evaluation result supports sample efficiency for GP-BO in the pulsed family. It does not establish universal BO superiority: random search outperformed the current GP setup for three-step charging, and the 80-evaluation overall difference was negligible. This should be presented as an optimizer-behaviour finding, not hidden.")
    add_paragraph(doc, "3.3 Physical interpretation of the hybrid objective", style="Heading 2")
    add_paragraph(doc, "In the present single-session runs, Q_cyclic dominates Q_calendar because calendar aging depends on a much longer time scale. Accordingly, the hybrid ranking mainly reflects the compromise between charge time and cycling stress. A fast pulsed profile can rank first even with a slightly higher cyclic-loss index, because the time term is lower. This is a property of the stated weights, not evidence that pulsing universally extends battery life.")
    add_note(doc, "Interpretation boundary", "Do not claim lithium plating, SEI thickness, resistance growth, LLI, or LAM has been directly measured. The current model uses stress-informed semi-empirical capacity-loss relations. High SoC, low temperature and high C-rate may be presented only as plating-risk proxies unless a negative-electrode-potential or validated electrochemical model is added.")
    add_paragraph(doc, "3.4 Required figures for the supervisory review", style="Heading 2")
    add_table(doc, ["Figure", "Axes / variables", "Physical question answered"], [
        ["Calendar-loss contour", "mean SoC x T, colour Q_calendar", "How SoC and temperature jointly accelerate calendar aging"],
        ["Cyclic-loss curves", "Ah throughput x Q_cyclic; C/2, 2C, 6C, 10C", "How the Eq. 7/Table 7 model responds to C-rate"],
        ["Cumulative degradation trajectory", "time; I, V, T, SoC, cumulative Q_calendar and Q_cyclic", "When each profile accumulates stress during charging"],
        ["Equal-energy profile table", "time, peak T, mean C-rate, Ah, Qcal, Qcyc", "Which physical trade-off determines the selected family"],
        ["Ambient sensitivity", "15, 25, 35 deg C; time, Tpeak, Qcal, Qcyc", "Whether the selected strategy changes under thermal stress"],
        ["Thermal cross-check", "BDT Tpeak versus lumped-model Tpeak", "Magnitude of thermal-model disagreement"],
    ], [2000, 3350, 4010])

    add_paragraph(doc, "4. Implementation plan before Tuesday", style="Heading 1")
    add_number(doc, "Audit the equations and units. Confirm that the primary run uses Eq. 2 for calendar loss and Eq. 7/Table 7 for cyclic loss. Record coefficient source, chemistry, temperature unit, time unit, Ah unit, and whether Q_loss is normalized.")
    add_number(doc, "Add EFC, mean SoC and DeltaSoC to every profile report. These are needed to connect the energy-constrained session to Eq. 4 and depth-of-discharge stress.")
    add_number(doc, "Generate the six physical figures listed in Section 3.4 from the same fixed scenario. Do not change energy targets or constraints across compared profiles.")
    add_number(doc, "Run one reproducible RW9 hybrid comparison: GP-BO and random search, same seed and 40 evaluations per family. Preserve JSON, figure, command line, and Git commit in a run manifest.")
    add_number(doc, "Run an ambient sweep at 15, 25 and 35 deg C for at least CCCV, two-step and pulsed profiles. Report Q_calendar and Q_cyclic separately.")
    add_number(doc, "Implement a structured research dashboard only after the model report is stable. Inputs are numeric fields/dropdowns; no natural-language model is required.")
    add_number(doc, "Create the RL environment after the hybrid step reward has been verified. Do not present trained RL performance until it is compared with the same BO and CC baselines.")

    add_paragraph(doc, "5. Structured daily-use interface specification", style="Heading 1")
    add_paragraph(doc, "The interface should be a reproducible experiment and decision dashboard, not a natural-language assistant. It converts explicit user/system inputs into a constrained charging request. The interface does not itself establish novelty; the research contribution is the goal-conditioned, degradation-aware controller beneath it.")
    add_table(doc, ["Input", "Unit / control", "Physical role"], [
        ["Cell ID / chemistry", "dropdown", "Selects capacity, OCV, voltage and current limits"],
        ["Initial SoC", "%", "Initial lithium inventory / energy state"],
        ["Initial and ambient temperature", "deg C", "Electro-thermal state and aging stress"],
        ["Relative age / SoH", "0-1 or %", "Capacity and resistance state"],
        ["Required energy", "Wh or pack-energy fraction", "Hard service requirement"],
        ["Available charging time", "min", "Deadline constraint"],
        ["Charger current limit", "A", "C-rate and power bound"],
        ["Maximum cell temperature", "deg C", "Thermal safety limit"],
        ["Objective mode", "fastest / balanced / minimum loss", "Selects documented reward weights"],
    ], [2600, 2050, 4710])
    add_paragraph(doc, "The outputs should be the top three feasible profiles, each with a current schedule, delivered energy, completion time, peak voltage, peak temperature, average C-rate, Ah throughput, Q_calendar, Q_cyclic, total loss index, and explicit constraint margins.")

    add_paragraph(doc, "6. Phase prompts for future Codex use", style="Heading 1")
    prompt1 = '''You are working in the battery_RW_NASA repository. Implement Phase 1: a physically interpretable hybrid degradation analysis for energy-constrained charging. Do not introduce an LLM or change unrelated code. First inspect Constrained_BO/hybrid_degradation.py, objective.py, simulator.py, run.py, and the existing RW9 result JSON files. Treat Eq. 2 (SoC-T-time Arrhenius) as the primary calendar model and Eq. 7/Table 7 (C-rate-T-Ah throughput) as the primary cyclic model. Eq. 3 and Eq. 5 must be selectable calendar sensitivity alternatives; Eq. 8 must be clearly marked as an alternative cyclic model unless its chemistry-specific coefficients are calibrated. Do not sum alternative equations. Add explicit model metadata and units to exported results; call the output a relative capacity-loss index unless a calibration establishes percent capacity fade. Add EFC, mean SoC, DeltaSoC, Ah throughput, mean/max C-rate, energy delivered/required, peak V, peak T, and constraint margins to per-profile metrics. Add deterministic tests for monotonic behavior: calendar loss increases with SoC/T/time; cyclic loss increases with Ah and C-rate under the selected model. Add a plotting/report script that generates: (1) Eq. 2 SoC-T contour, (2) Eq. 7/Table 7 Qcyc-vs-Ah curves at C/2, 2C, 6C, 10C, (3) cumulative Qcal/Qcyc along each best profile, and (4) equal-energy profile comparison table. Preserve backward compatibility, update README methodology wording, run the tests, and report exact files changed plus any limitations due to chemistry-specific coefficients.'''
    add_code_prompt(doc, "Prompt 1 - Phase 1: physically interpretable hybrid degradation model", prompt1)
    prompt2 = '''You are working in the battery_RW_NASA repository. Implement Phase 2: a safe, goal-conditioned reinforcement-learning environment on top of the verified hybrid-degradation charging simulator. This is not a natural-language interface. Build a reproducible Python environment whose observation at each decision interval includes SoC, terminal voltage, cell temperature, relative age/SoH, previous current, remaining energy requirement, remaining time, ambient temperature, and optional uncertainty margin. Start with a discrete action set of safe charging-current levels; do not add deep-RL dependencies until the environment and tests work. Use the existing FrozenBDT/ChargingSimulator rollout and compute_step_reward so the per-step reward is incremental delivered energy minus Qcalendar increment, Qcyclic increment, and time increment. Enforce hard safety by action masking/projection before rollout: V <= Vmax, T <= Tmax, |I| <= Imax, and session must meet the remaining-energy/deadline requirement. Add a Gymnasium-compatible wrapper only if Gymnasium is already available; otherwise implement a small dependency-free environment with reset/step. Add deterministic tests for energy accounting, action rejection, reward decomposition, terminal conditions, and reproducibility. Provide baseline evaluation hooks for CCCV, the saved BO profile, and a simple tabular Q-learning agent. Save each episode with physical metrics (Wh, Ah, C-rate, Qcalendar, Qcyclic, peak V/T, feasibility), not just reward. Do not claim RL results until the agent is trained across multiple seeds and evaluated under exactly the same initial-state and energy/deadline scenarios as BO.'''
    add_code_prompt(doc, "Prompt 2 - Phase 2: safe goal-conditioned RL environment", prompt2)

    add_paragraph(doc, "7. References used in this draft", style="Heading 1")
    add_paragraph(doc, "Sinha, S. S., Lehman, B., and Bharadwaj, P. (2026). Life Extension of Lithium-Ion Battery: Degradation Comprehension, Modeling, Characterization, and Mitigation Strategies. IEEE Open Journal of Power Electronics, 7. DOI: 10.1109/OJPEL.2025.3639205.")
    add_paragraph(doc, "Repository evidence: Constrained_BO/results/GP_BO_vs_random_search_RW9.md; RW9 and RW9_random constrained_bo_results_n_calls_40.json and n_calls_80.json; outputs/finetune_two_stage_RW10-RW12/finetune_percent_summary.json.", size=10, color=GRAY)

    doc.save(OUT)
    print(OUT)

if __name__ == "__main__":
    main()
